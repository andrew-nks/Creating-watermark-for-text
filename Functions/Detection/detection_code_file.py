# detect encoded characters
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def read_encoded_characters_from_code_file(file_path,output_path):
    
    # Initialize an empty string to store each line of code
    current_line = ''

    with open(file_path, 'r', encoding='utf-8') as file:

        # count of non-encoded characters
        score = 0

        # count of encoded characters
        encoded = 0

        whitespace_character = "\u200B"  # Four-Per-Em Space character

        # count of whitespaces
        whitespace = 0

        # homoglyphs
        homoglyphs = "АаВеցіΚӏΜΝոΟΡрԛЅѕΤՍԜԝΥу‚;꞉ǃʾ"

        # read all lines in code
        lines = file.readlines()


        # create new document with flagged homoglyphs
        new_doc = Document()

        for line in lines:
            paragraph = new_doc.add_paragraph()
            for letter in line:
                run = paragraph.add_run(letter)
                # Check if the letter is in the flag list
                if letter in homoglyphs:
                    # Apply yellow highlighting
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
             
        
        # save doc
        new_doc.save(output_path)
        

        # Flag to track if we are currently inside a comment block
        inside_comment = False

        