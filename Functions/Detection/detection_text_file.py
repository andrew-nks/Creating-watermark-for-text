# detect encoded punctuation
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def read_encoded_characters_from_word_file_with_paragraphs(file_path, output_path):
    whitespace_character = "\u200B"  # Four-Per-Em Space character

    # count of non-encoded characters
    score = 0

    # count of encoded characters
    encoded = 0

    # count of whitespaces
    whitespace = 0

    # count of words
    word_count = 0

    # homoglyphs
    homoglyphs = "АаВеցіΚӏΜΝոΟΡрԛЅѕΤՍԜԝΥу‚;꞉ǃʾ"

    # Initialize an empty list to store the words
    words = []

    try:
        # Open the Word document
        doc = Document(file_path)

        # create new document with flagged homoglyphs
        new_doc = Document()

        # Iterate through each paragraph in the document
        for paragraph in doc.paragraphs:
            # Add the original paragraph text to the list
            words.extend(paragraph.text.split())

            new_para = new_doc.add_paragraph()  # Add a new paragraph to the new document

            # Iterate through characters in the paragraph
            for char in paragraph.text:
                run = new_para.add_run(char)  # Add a run for each character
                
                # Check if the character should be flagged
                if char in homoglyphs:
                    font = run.font
                    font.highlight_color = WD_COLOR_INDEX.YELLOW 
        
            # Add a new line after each paragraph
            new_para.add_run('\n')
    
        # Save the new document
        new_doc.save(output_path)
        
    except Exception as e:
        # Handle any exceptions that occur during file processing
        print(f"Error reading file: {e}")
    
    
    
   

    # iterate through each word and each character in text file
    for word in words:
        for ch in word:
            if ch in homoglyphs:
                encoded += 1
            elif ch == whitespace_character:
                whitespace += 1
            else:
                score += 1
        word_count += 1

    

    proportion_of_encoded_characters = str(round(encoded/(encoded + score),4))

    if whitespace > word_count:
         proportion_of_whitespace = str(round(whitespace/score,4))     # whitespace is after every character
    else:
         proportion_of_whitespace = str(round(whitespace/word_count,4))  # whitespace is after every word

    return [proportion_of_encoded_characters, proportion_of_whitespace]
