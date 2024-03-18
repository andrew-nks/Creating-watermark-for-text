from docx import Document

def read_words_from_word_file_with_paragraphs(file_path):
    
    # Initialize an empty list to store the paragraphs
    paragraphs = []

    try:
        # Open the Word document
        doc = Document(file_path)

        # Iterate through each paragraph in the document
        for paragraph in doc.paragraphs:
            # Add the original paragraph text to the list
            paragraphs.append(paragraph.text)

        
    except Exception as e:
        # Handle any exceptions that occur during file processing
        print(f"Error reading file: {e}")
    
    return paragraphs

def write_words_to_word_file_with_paragraphs(paragraphs, output_file_path):
    try:
        # Create a new Word document
        doc = Document()

        

        # Add each paragraph to the document
        for paragraph_string in paragraphs:
            # split each paragraph string into a list of word strings
            paragraph_list_of_words = paragraph_string.split()

            # new list of encoded words per paragraph list
            new_paragraph_list_of_words = []

            # iterate through each list of word strings to check each word for punctuation
            for i in range(len(paragraph_list_of_words)):
                word = paragraph_list_of_words[i]
                
                if word[-1] == ",":
                    new_paragraph_list_of_words.append(word[0:len(word)-1]+ "‚"+" ")   # this is unrecognizable
                elif word[-1] == ";":
                    new_paragraph_list_of_words.append(word[0:len(word)-1]+";"+" ")   # this is unrecognizable
                elif word[-1] == ":":
                    new_paragraph_list_of_words.append(word[0:len(word)-1]+"꞉"+" ")   # this is unrecognizable
                elif word[-1] == "*":
                    new_paragraph_list_of_words.append(word[0:len(word)-1]+"⃰"+" ")   # this appears different
                elif word[-1] == "'" or word[-1] == "‘" or word[-1] == "’":
                    new_paragraph_list_of_words.append(word[0:len(word)-1]+"῾"+" ")   # this appears different
    
                elif i == len(paragraph_list_of_words) - 1:                           # no white space after last word of paragraph
                    new_paragraph_list_of_words.append(word)    
                else:
                    new_paragraph_list_of_words.append(word + " ") # every word that is not last must have whitespace
            
            new_paragraph_string = ''.join(new_paragraph_list_of_words)
            
            doc.add_paragraph(new_paragraph_string)

        # Save the document
        doc.save(output_file_path)

        print(f"Paragraphs written to {output_file_path} successfully.")
    except Exception as e:
        # Handle any exceptions that occur during file writing
        print(f"Error writing file: {e}")

       


